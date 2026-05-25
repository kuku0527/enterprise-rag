"""Document Loaders for various file formats"""

import os
from pathlib import Path
from typing import List, Optional
from abc import ABC, abstractmethod

from loguru import logger


class Document:
    """Simple document representation"""

    def __init__(self, content: str, metadata: Optional[dict] = None):
        self.content = content
        self.metadata = metadata or {}

    def __repr__(self):
        return f"Document(content={self.content[:50]}..., metadata={self.metadata})"


class BaseLoader(ABC):
    """Base loader interface"""

    @abstractmethod
    def load(self, file_path: str) -> List[Document]:
        pass

    @abstractmethod
    def load_dir(self, dir_path: str, glob_pattern: str = "**/*") -> List[Document]:
        pass


class TextLoader(BaseLoader):
    """Load plain text files"""

    def load(self, file_path: str) -> List[Document]:
        logger.info(f"Loading text file: {file_path}")
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        metadata = {
            "source": file_path,
            "file_type": "txt",
            "file_name": os.path.basename(file_path)
        }
        return [Document(content=content, metadata=metadata)]

    def load_dir(self, dir_path: str, glob_pattern: str = "**/*.txt") -> List[Document]:
        documents = []
        dir_path = Path(dir_path)

        for file_path in dir_path.glob(glob_pattern):
            if file_path.is_file():
                try:
                    docs = self.load(str(file_path))
                    documents.extend(docs)
                except Exception as e:
                    logger.error(f"Error loading {file_path}: {e}")

        logger.info(f"Loaded {len(documents)} documents from {dir_path}")
        return documents


class PDFLoader(BaseLoader):
    """Load PDF files"""

    def load(self, file_path: str) -> List[Document]:
        logger.info(f"Loading PDF file: {file_path}")

        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("Please install pypdf: pip install pypdf")

        reader = PdfReader(file_path)
        documents = []

        for i, page in enumerate(reader.pages):
            content = page.extract_text()
            if content.strip():
                metadata = {
                    "source": file_path,
                    "file_type": "pdf",
                    "file_name": os.path.basename(file_path),
                    "page_number": i + 1
                }
                documents.append(Document(content=content, metadata=metadata))

        logger.info(f"Extracted {len(documents)} pages from PDF")
        return documents

    def load_dir(self, dir_path: str, glob_pattern: str = "**/*.pdf") -> List[Document]:
        documents = []
        dir_path = Path(dir_path)

        for file_path in dir_path.glob(glob_pattern):
            if file_path.is_file():
                try:
                    docs = self.load(str(file_path))
                    documents.extend(docs)
                except Exception as e:
                    logger.error(f"Error loading {file_path}: {e}")

        logger.info(f"Loaded {len(documents)} PDF documents from {dir_path}")
        return documents


class DocxLoader(BaseLoader):
    """Load Word documents"""

    def load(self, file_path: str) -> List[Document]:
        logger.info(f"Loading DOCX file: {file_path}")

        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Please install python-docx: pip install python-docx")

        doc = DocxDocument(file_path)
        content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        metadata = {
            "source": file_path,
            "file_type": "docx",
            "file_name": os.path.basename(file_path)
        }
        return [Document(content=content, metadata=metadata)]

    def load_dir(self, dir_path: str, glob_pattern: str = "**/*.docx") -> List[Document]:
        documents = []
        dir_path = Path(dir_path)

        for file_path in dir_path.glob(glob_pattern):
            if file_path.is_file():
                try:
                    docs = self.load(str(file_path))
                    documents.extend(docs)
                except Exception as e:
                    logger.error(f"Error loading {file_path}: {e}")

        logger.info(f"Loaded {len(documents)} DOCX documents from {dir_path}")
        return documents


class DocumentLoader:
    """Unified document loader factory"""

    LOADER_MAP = {
        ".txt": TextLoader,
        ".pdf": PDFLoader,
        ".docx": DocxLoader,
        ".doc": DocxLoader,
    }

    @classmethod
    def get_loader(cls, file_path: str) -> BaseLoader:
        """Get appropriate loader based on file extension"""
        ext = Path(file_path).suffix.lower()
        loader_class = cls.LOADER_MAP.get(ext)

        if not loader_class:
            raise ValueError(f"Unsupported file type: {ext}")

        return loader_class()

    @classmethod
    def load_file(cls, file_path: str) -> List[Document]:
        """Load a single file"""
        loader = cls.get_loader(file_path)
        return loader.load(file_path)

    @classmethod
    def load_directory(cls, dir_path: str) -> List[Document]:
        """Load all supported files from a directory"""
        documents = []
        dir_path = Path(dir_path)

        for ext, loader_class in cls.LOADER_MAP.items():
            loader = loader_class()
            docs = loader.load_dir(str(dir_path), f"**/*{ext}")
            documents.extend(docs)

        logger.info(f"Total documents loaded from {dir_path}: {len(documents)}")
        return documents
